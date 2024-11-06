# helpers.py

import os
import pandas as pd
import re
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from nltk import ngrams
import csv
import fitz
from flask import flash

users = [
    {"id": 1, "username": "john_doe", "email": "john@example.com", "role": "Student"},
    {"id": 2, "username": "jane_smith", "email": "jane@example.com", "role": "Lecturer"}
]
posts = [
    {"id": 1, "title": "Sample Post 1", "content": "This is the first sample post.", "user_id": 1, "comments": [], "category": "General"},
    {"id": 2, "title": "Sample Post 2", "content": "This is the second sample post.", "user_id": 2, "comments": [
        {"content": "This is a comment on post 2", "user_id": 1}
    ], "category": "Academic"}
]

def get_post_by_id(post_id, posts):
    return next((post for post in posts if post['id'] == post_id), None)

def get_comments_for_post(post_id, comments):
    return [comment for comment in comments if comment['post_id'] == post_id]

def get_user(user_id, users):
    for user in users:
        if user['id'] == user_id:
            return user
    return None

def read_job_data(file_path):
    df = pd.read_excel(file_path)
    df.fillna('', inplace=True)
    return df

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() not in full_text:
                    full_text.append(cell.text.strip())

    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                full_text.append(para.text.strip())
        if section.footer:
            for para in section.footer.paragraphs:
                full_text.append(para.text.strip())

    return "\n".join([text for text in full_text if text])

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        flash('Error reading the PDF file. Please upload a valid resume.', 'danger')
        return ""
    return text

def preprocess_text(text):
    stop_words = set(stopwords.words('english'))
    text = re.sub(r'\b\d{1,4}\b', '', text)
    text = re.sub(r'\W', ' ', text)
    words = [word for word in text.lower().split() if word not in stop_words]
    bigrams = [' '.join(gram) for gram in ngrams(words, 2)]
    trigrams = [' '.join(gram) for gram in ngrams(words, 3)]
    return set(words + bigrams + trigrams)

def calculate_similarity(resume_text, job_data):
    corpus = [resume_text] + job_data['combined_text'].tolist()
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(corpus)
    cosine_similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    recommendations = list(zip(job_data['occupation'], cosine_similarities))
    recommendations.sort(key=lambda x: x[1], reverse=True)
    return [job for job, score in recommendations]

def load_club_data():
    clubs = []
    with open('static/data/clubs_data.csv', newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            clubs.append({
                'name': row['Club_Name'].strip(),
                'logo': row['Logo_URL'].strip() if row['Logo_URL'].strip() else None,
                'email': row['Mail_ID'].strip(),
                'website': row['Website_Link'].strip()
            })
    return clubs
